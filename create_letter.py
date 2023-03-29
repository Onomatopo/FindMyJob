import pandas as pd
import openai
import docx
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from datetime import datetime
import re
from docx2pdf import convert
import win32com.client
import os
from data_extraction import cookies_ablehnen, chromedriver_aufrufen
from personal_info import name


# ==============================
# Hauptfunktionen
# ==============================

# diese Funktion erstellt und speichert das gesamte Anschreiben
def anschreiben_erstellen(df, api_key, eigene_adresse, fruehestes_datum, unterschrift, schwellenwerte, anschreiben_speicherpfad, anzeigen_speicherpfad):
    # Ähnlichkeitsschwellenwerte einlesen:
    aufgaben_schwellenwert = schwellenwerte[0]
    profil_schwellenwert = schwellenwerte[1]
    benefits_schwellenwert = schwellenwerte[2]
    gesamt_schwellenwert = schwellenwerte[3]

    # konstantere (Text)bausteine des Anschreiben
    fußnote_1 = 'https://chat.openai.com/'
    fußnote_2 = 'https://github.com/Onomatopo/MyJobFinder.git'
    heute = datetime.today().strftime('%d.%m.%Y') # brauchen wir später zum einfügen von ort und datum im anschreiben

    anschreiben_outro = "Als Teamplayer schätze ich die Zusammenarbeit mit Kollegen und bringe mich gerne in " \
                        "interdisziplinäre Teams ein. Meine analytische Denkweise und meine Begeisterung für Daten " \
                        "ermöglichen es mir, komplexe Probleme zu verstehen und effektive Lösungen zu entwickeln. In " \
                        "verschiedenen eigenen Projekten konnte ich meine Fähigkeiten in Datenanalyse und -verarbeitung " \
                        "erfolgreich anwenden. So ist beispielsweise diese Bewerbung automatisiert durch Webscraping von " \
                        "Jobanzeigen, Ansätzen des Natural Language Processing (NLP) zur Verarbeitung und inhaltlichen " \
                        "Bewertung des Textes entstanden. Zudem sind Teile dieses Anschreibens von ChatGPT¹ verfasst und " \
                        "per API-Schnittstelle hier eingefügt worden. (Der zugrundeliegende Code ist auf meinem " \
                        "GitHub-Profil² einsehbar.)\nGerne würde ich meine Erfahrungen und Kenntnisse im Rahmen Ihrer " \
                        "anspruchsvollen Aufgabenstellung einbringen und mich gemeinsam mit Ihrem Team weiterentwickeln. " \
                        "Aufgrund meiner bisherigen Erfahrungen bin ich davon überzeugt, dass ich die Anforderungen der " \
                        "ausgeschriebenen Stelle erfüllen und zum Erfolg des Unternehmens beitragen kann. Mein " \
                        f"frühestmögliches Eintrittsdatum is der {fruehestes_datum}.\nIch bedanke mich für Ihre " \
                        f"Aufmerksamkeit und freue mich sehr auf Ihre Rückmeldung." \
                        "\n\nMit freundlichen Grüßen," \

    # variable Textbausteine werden direkt in der loop zur erstellung der anschreiben eingefügt
    for index, row in df.iterrows():
        anschreiben_intro = f'mit großem Interesse habe ich Ihre Stellenausschreibung für ' \
                            f'einen {df.loc[index, "Jobtitel"]} gelesen. Als Absolvent der Wirtschaftswissenschaften ' \
                            f'und derzeitiger Mitarbeiter im Prozessmanagement bringe ich umfangreiche Kenntnisse und ' \
                            f'Erfahrungen in der Analyse von Geschäfts- und IT-prozessen mit. Nun strebe ich an, meine Fähigkeiten ' \
                            f'und Kompetenzen in den Bereich der Datenanalyse zu erweitern und bewerbe mich daher bei Ihnen.'

        # bedingungen für das erstellen des anschreibens festlegen (mindest-"ähnlichkeitswert")
        # (Ähnlichkeitswerte für Aufgaben, Profil, Benefits jeweils mind. 20) oder Gesamtähnlichkeit von mind. 60 und sofern noch kein anschreiben erstellt wurde
        if ((df.loc[index, 'aufgaben_score_jaccard'] > aufgaben_schwellenwert) and (df.loc[index, 'profil_score_jaccard'] > profil_schwellenwert) and (df.loc[index, 'benefits_score_jaccard'] > benefits_schwellenwert) and ((df.loc[index, 'anschreiben_erstellt?'] != "Ja") and (df.loc[index, 'anschreiben_erstellt?'] != "Ja"))) or ((df.loc[index, 'Gesamtscore_jaccard'] > gesamt_schwellenwert) and (df.loc[index, 'anschreiben_erstellt?'] != "Ja")):

            # neues word document pro bewerbung erstellen
            doc = docx.Document()
            # Textstyles definieren
            text_normal = doc.styles.add_style('text_normal', WD_STYLE_TYPE.PARAGRAPH)
            text_normal.font.name = 'Calibri'
            text_normal.font.size = Pt(12)
            fußnoten_style = doc.styles.add_style('fußnote', WD_STYLE_TYPE.PARAGRAPH)
            fußnoten_style.font.name = 'Calibri'
            fußnoten_style.font.size = Pt(8)
            fett_style = doc.styles.add_style('fett', WD_STYLE_TYPE.PARAGRAPH)
            fett_style.font.name = 'Calibri'
            fett_style.font.size = Pt(12)
            fett_style.font.bold = True

            # Anschreiben Text erstellen
            # eigenen adressblock und ort + datum einfügen und rechts ausrichten
            mein_adressblock = doc.add_paragraph(f"{eigene_adresse}", style='text_normal')
            mein_adressblock.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            ort_datum = doc.add_paragraph(style='text_normal')
            ort_datum.add_run(f"Hamburg, den {heute}")
            ort_datum.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

            # "header" (firmen name hinzufügen)
            header = doc.add_paragraph(df.loc[index, "Firma"], style='fett')

            # Adresse nur in brief einfügen falls das feld im DF gefüllt ist
            if not pd.isnull(df.loc[index, 'Adresse']):
                # ziel adresse (muss erstmal überarbeitet werden, damit es zu einer liste wird)
                ziel_adresse_liste = [part.strip(" '") for part in df.loc[index, "Adresse"].strip("[]").split(",")]
                ziel_adresse = ""
                for zeile in ziel_adresse_liste:
                    ziel_adresse += f"{zeile}\n"
                doc.add_paragraph(f'{ziel_adresse}', style='text_normal')
            else:
                # falls keine Adresse hinterlegt --> 2 Leerzeilen einfügen
                doc.add_paragraph("\n\n", style='text_normal')

            # "überschrift" einfügen
            doc.add_paragraph(f'Bewerbung um die Stelle als {df.loc[index, "Jobtitel"]}', style='fett')

            # intro paragraph einfügen
            doc.add_paragraph('Sehr geehrte Damen und Herren,', style='text_normal')
            doc.add_paragraph(anschreiben_intro, style='text_normal')

            # den anschreiben_mid nur erstellen falls die Schwellenwerte erreicht (daher in dieser loop)
            # # hier API call zu ChatGPT
            openai.api_key = api_key
            chat_prompt = f"create a text in german which tasks appeal to me in particular as {df.loc[index, 'Jobtitel']}" \
                          f"(not female perspective)" \
                          f"and why the company {df.loc[index, 'Firma']} itself is particular appealing to me in not more " \
                          f"than 130 words: {df.loc[index, 'matched_tasks']}"
            # Fehlerhandling der API Connection handlen
            try:
                completion = openai.ChatCompletion.create(model="gpt-3.5-turbo",
                                                          messages=[{"role": "user", "content": chat_prompt}])
                # Ergebnis in DF speichern
                df.loc[index, 'anschreiben_mid'] = completion.choices[0].message['content']
                print(f'prompt {index} fertig')
            except openai.error.APIConnectionError:
                df['anschreiben_mid'] = ""
            # Leerzeilen aus dem Abschnitt entfernen
            df.loc[index, 'anschreiben_mid'] = df.loc[index, 'anschreiben_mid'].replace('\n\n', ' ')
            df.loc[index, 'anschreiben_mid'] = df.loc[index, 'anschreiben_mid'].strip()

            # anschreiben_mid nur in brief einfügen falls das feld im DF gefüllt ist
            if not pd.isnull(df.loc[index, 'anschreiben_mid']):
                # mid paragraph einfügen
                doc.add_paragraph(df.loc[index, 'anschreiben_mid'], style='text_normal')

            # outro paragraph einfügen
            doc.add_paragraph(anschreiben_outro, style='text_normal')

            # Bild der Unterschrift einfügen
            doc.add_picture(unterschrift)
            # (Paul Schmidt) einfügen
            doc.add_paragraph("(Paul Schmidt)", style='text_normal')

            # Fußnoten einfügen, workaround via Fußzeilen und Einstellung, dass Fußzeile auf Seite 1 unterschiedlich ist
            # Abschnitt definieren + Fußzeilen einfügen
            fußzeile_abschnitt = doc.sections[0]
            fußzeile = fußzeile_abschnitt.footer
            fußzeile.add_paragraph(f"¹ {fußnote_1}\n² {fußnote_2}", style='fußnote')
            # Einstellung, dass die 1. Seite eine eigene Fußzeile hat
            fußzeile_abschnitt.different_first_page_header_footer = True

            # anschreiben als docx speichern
            dateiname = f'Bewerbung_{df.loc[index, "Firma"]}_{df.loc[index, "Jobtitel"]}'
            dateiname = re.sub(r'[/:<>*?"\\|]', '_', str(dateiname))
            doc.save(f'{anschreiben_speicherpfad}/{dateiname}.docx')
            # das anschreiben auch als pdf speichern
            convert(f'{anschreiben_speicherpfad}/{dateiname}.docx', f'{anschreiben_speicherpfad}/{dateiname}.pdf')
            # dateinamen in spalte speichern (für später um als anhang in email zu öffnen)
            df.loc[index, "anschreiben_dateiname"] = dateiname

            # ort_datum clearen, weil das sonst dazu führen kann dass es nicht richtig eingefügt wird
            ort_datum.clear()

            # im DF markieren, dass Anschreiben erstellt wurde
            df.loc[index, "anschreiben_erstellt?"] = "Ja"

            # die Jobanzeige speichern
            # (dazu müssen die anzeigen nochmal per driver aufgerufen werden)
            driver = chromedriver_aufrufen()
            driver.get(df.loc[index, "URL"])
            cookies_ablehnen(driver)
            html = driver.page_source
            # die html seite abspeichern (unerlaubte zeichen per regex entfernen)
            dateiname = df.loc[index, "Firma"] + '_' + df.loc[index, "Jobtitel"]
            dateiname = re.sub(r'[/:<>*?"\\|]', '_', str(dateiname))
            with open(f'{anzeigen_speicherpfad}/{dateiname}.html', 'w', encoding='utf-8') as f:
                f.write(html)
            driver.quit()

    return df

# diese funktion erstellt einen entwurf der mail mit dem anschreiben, lebenslauf, zeugnisse und kurzem text
def mail_erstellen(df, anschreiben_speicherpfad, lebenslauf_speicherpfad, zeugnisse_speicherpfad):
    import win32com.client as win32
    # eine instanz von outlook erstellen
    outlook = win32com.client.Dispatch('Outlook.Application')
    # wie viele e-mails maximal geöffnet werden können (anzahl der zeilen im DF)
    # das ergebnis von .shape ist (x, x) für (zeilen, spalten) --> daher [0]
    anzahl_zeilen = df.shape[0]

    # loop über alle zeilen des DF
    for i in range(anzahl_zeilen):
        # prüfen ob eine e-Mail hinterlegt wurde & das anschreiben erstellt wurde & und die mail noch nicht verschickt bzw gedraftet wurde
        if (df.loc[i, 'email'] is not pd.NA) and (df.loc[i, 'anschreiben_erstellt?'] == "Ja") and (df.loc[i, 'bewerbung_verschickt?'] != "Ja"):
            mail = outlook.CreateItem(0)
            # empfänger, betreff, text und anhänge der mail
            mail.To = df['email'][i]
            mail.Subject = f'Bewerbung als {df.loc[i, "Jobtitel"]}'
            mail.Body = f'Sehr geehrte Damen und Herren,\nbitte entnehmen Sie dem Anhang meine Bewerbung um die ' \
                            f'ausgeschriebene Stelle als {df.loc[i, "Jobtitel"]}. Dort finden Sie in drei Dokumenten das ' \
                            f'Anschreiben, meinen Lebenslauf sowie die Arbeitgeberzeugnisse.' \
                            f'\nmit freundlichen Grüßen,\n{name}'

            # anhänge beifügen (anschreiben + lebenslauf + zeugnisse)
            anhang_anschreiben = f"{anschreiben_speicherpfad}/{df.loc[i, 'anschreiben_dateiname']}.pdf"
            anhang_lebenslauf = lebenslauf_speicherpfad
            anhang_zeugnisse = zeugnisse_speicherpfad
            mail.Attachments.Add(os.path.join(os.getcwd(), anhang_anschreiben))
            mail.Attachments.Add(os.path.join(os.getcwd(), anhang_lebenslauf))
            mail.Attachments.Add(os.path.join(os.getcwd(), anhang_zeugnisse))

            # spalte als verschickt markieren
            df.loc[i, 'bewerbung_verschickt?'] = "Ja"

            # email im entwürfe ordner speichern
            mail.Save()
